"""Fill the real Custom.docx / Ready to Go.docx templates with
proposal data.

Rationale — the HTML template shipped in the first pass was a visual
approximation of the Word document. R&D / sales want the *exact*
file they already use as a contract deliverable, with the Vita NPD
letterhead, fonts, and spacing preserved. Re-implementing that in
HTML would mean sweating every pixel; python-docx on the original
files is strictly faster and always matches.

Two render modes:

* ``render_docx_bytes`` — returns the filled .docx so sales can
  download, tweak manually if needed, and email to the client.
* ``render_pdf_bytes`` — converts the filled .docx to PDF via
  :mod:`docx2pdf` (drives Microsoft Word on macOS) so the browser /
  kiosk iframe can display it inline. Falls back to ``None`` when
  Word / LibreOffice are not installed; the caller picks up the
  HTML fallback.
"""

from __future__ import annotations

import base64
import copy
import io
import os
import re
import shutil
import subprocess
import tempfile
import threading
import zipfile
from dataclasses import dataclass
from decimal import Decimal
from pathlib import Path
from typing import Iterable

from lxml import etree


#: Global lock that serialises external-converter calls. LibreOffice
#: tolerates parallel invocations only when each gets its own user
#: profile, but a single dev machine still benefits from serialising
#: — two parallel iframe loads + a download click would otherwise
#: spawn three soffice processes fighting for the same CPU core.
_RENDER_LOCK = threading.Lock()

#: Bump this integer whenever the template filler, formatting rules
#: or post-processing change in a way that affects the pixel output.
#: It's mixed into :func:`_proposal_digest` so every cached PDF gets
#: regenerated after a deploy even when the underlying proposal
#: hasn't been edited. History:
#:   1 → 2: £/€/$ glyph prefix on prices
#:   2 → 3: drop the "Vita NPD Sales Team" company line from signoff
#:   3 → 4: proposal-level sales_person override used in signoff
_RENDER_VERSION = 4

from django.conf import settings
from docx import Document
from docx.document import Document as _Document
from docx.shared import Inches
from docx.table import _Cell, Table

from apps.proposals.models import Proposal, ProposalTemplateType


TEMPLATE_DIR = (
    Path(settings.BASE_DIR) / "apps" / "proposals" / "templates" / "proposals"
)
TEMPLATE_PATHS: dict[str, Path] = {
    ProposalTemplateType.CUSTOM.value: TEMPLATE_DIR / "custom_template.docx",
    ProposalTemplateType.READY_TO_GO.value: TEMPLATE_DIR / "ready_to_go_template.docx",
}

#: Hints for the "Dear ," greeting and the "Ref: – -" line. These
#: literal strings appear in both templates; we find them and splice
#: the real values in without disturbing the surrounding formatting.
_GREETING_PATTERNS = (
    "Dear ,",
    "Dear  ,",
)
_REF_PATTERNS = (
    "Ref:  –  - ",
    "Ref:  –  -",
    "Ref: –  - ",
)


@dataclass
class ProposalLineRow:
    """One product row as we want it rendered in the pricing table."""

    product_code: str
    description: str
    quantity: str
    unit_price: str
    amount: str


@dataclass
class ProposalRenderContext:
    """Flat payload of every value we splice into the template.

    Kept as a dataclass so each rendering concern (paragraphs,
    customer table, product rows, signatures) can grab what it needs
    without re-reading off the ORM instance repeatedly.
    """

    reference: str
    dear_name: str
    customer_name: str
    customer_email: str
    customer_phone: str
    customer_company: str
    invoice_address: str
    delivery_address: str
    product_code: str
    product_description: str
    quantity: str
    unit_price: str
    amount: str
    #: One row per :class:`ProposalLine`. Always at least one entry
    #: after the data migration backfills; multi-product proposals
    #: stack more rows here. The pricing table grows rows to match.
    line_rows: list[ProposalLineRow]
    lines_subtotal: str
    freight_amount: str
    total_excl_vat: str
    currency: str
    date_line: str
    #: Commercial owner whose name closes out the letter (replaces the
    #: hard-coded "Matthew Bowden" in the original templates). State
    #: transitions to ``in_review`` / ``approved`` / ``sent`` are gated
    #: on the project having a sales person assigned, so in practice
    #: this is always populated by the time a customer-facing render
    #: runs; draft previews may show an empty signature line.
    sales_person_name: str
    #: ``revisionnumber`` in the CRM template — pulled off the pinned
    #: formulation version so the header shows which snapshot the
    #: proposal was built against.
    revision_number: str
    #: Proposal creation date formatted for the header line.
    created_on: str
    customer_signature_image_png: bytes | None
    prepared_by_signature_image_png: bytes | None
    prepared_by_name: str
    prepared_by_date: str
    director_signature_image_png: bytes | None
    director_name: str
    director_date: str

    @classmethod
    def from_proposal(cls, proposal: Proposal) -> "ProposalRenderContext":
        version = proposal.formulation_version
        metadata = version.snapshot_metadata or {}
        currency = (proposal.currency or "GBP").strip() or "GBP"

        subtotal = proposal.subtotal
        total = proposal.total_excl_vat

        signed_at = proposal.customer_signed_at
        date_line = signed_at.strftime("%d %b %Y") if signed_at else ""

        # Sales person: the proposal-level override wins so multi-
        # project proposals (where the project's own assignment is
        # ambiguous) can nominate one signatory. Single-project
        # proposals leave the override null and fall back to the
        # formulation's ``sales_person`` — the scientist doesn't
        # re-pick for every proposal.
        formulation = version.formulation
        sales_person = proposal.sales_person or getattr(
            formulation, "sales_person", None
        )
        sales_name = ""
        if sales_person is not None:
            sales_name = (
                sales_person.get_full_name()
                or sales_person.email
                or ""
            ).strip()

        prepared_by = proposal.prepared_by_user
        director = proposal.director_user

        # Build a row per ProposalLine for the multi-product pricing
        # table. Lines are the source of truth post-migration; when
        # a proposal somehow has zero lines (shouldn't happen, but
        # guards against a half-migrated state) we fall back to the
        # Proposal-level fields so the render doesn't break.
        line_rows: list[ProposalLineRow] = []
        lines_total_decimal: Decimal | None = None
        lines = list(proposal.lines.select_related("formulation_version__formulation").order_by(
            "display_order", "created_at"
        ))
        if lines:
            for line in lines:
                row_subtotal = line.subtotal
                if row_subtotal is not None:
                    lines_total_decimal = (
                        row_subtotal
                        if lines_total_decimal is None
                        else lines_total_decimal + row_subtotal
                    )
                code_value = line.product_code or (
                    line.formulation_version.formulation.code
                    if line.formulation_version_id
                    else ""
                )
                desc_value = line.description or (
                    line.formulation_version.formulation.name
                    if line.formulation_version_id
                    else ""
                )
                line_rows.append(
                    ProposalLineRow(
                        product_code=code_value or "—",
                        description=desc_value or "—",
                        quantity=str(line.quantity or 1),
                        unit_price=_format_money(line.unit_price, currency),
                        amount=_format_money(row_subtotal, currency),
                    )
                )
        else:
            # Legacy single-product fallback — survives a proposal
            # saved between the migration and the line backfill.
            line_rows.append(
                ProposalLineRow(
                    product_code=(
                        metadata.get("code") or formulation.code or "—"
                    ),
                    description=(
                        metadata.get("name") or formulation.name or "—"
                    ),
                    quantity=str(proposal.quantity or 1),
                    unit_price=_format_money(proposal.unit_price, currency),
                    amount=_format_money(subtotal, currency),
                )
            )

        # Grand total includes freight when quoted. When every line
        # has a price, the lines_total_decimal is meaningful; when
        # any is missing we leave the footer as ``TBC`` so sales
        # notices before the proposal reaches a client.
        grand_total: Decimal | None = lines_total_decimal
        if grand_total is not None and proposal.freight_amount is not None:
            grand_total = grand_total + proposal.freight_amount
        # Prefer the multi-line total over the single-line envelope
        # total when any lines exist; keeps the two numbers in sync.
        total_for_template = grand_total if lines else total

        return cls(
            reference=proposal.reference or proposal.code or "",
            dear_name=(
                proposal.dear_name
                or proposal.customer_name
                or "Customer"
            ),
            customer_name=proposal.customer_name,
            customer_email=proposal.customer_email,
            customer_phone=proposal.customer_phone,
            customer_company=proposal.customer_company,
            invoice_address=proposal.invoice_address,
            delivery_address=proposal.delivery_address,
            product_code=line_rows[0].product_code if line_rows else "",
            product_description=(
                line_rows[0].description if line_rows else ""
            ),
            quantity=line_rows[0].quantity if line_rows else "1",
            unit_price=line_rows[0].unit_price if line_rows else "TBC",
            amount=line_rows[0].amount if line_rows else "TBC",
            line_rows=line_rows,
            lines_subtotal=_format_money(lines_total_decimal, currency),
            freight_amount=_format_money(proposal.freight_amount, currency),
            total_excl_vat=_format_money(total_for_template, currency),
            currency=currency,
            date_line=date_line,
            sales_person_name=sales_name,
            revision_number=str(version.version_number) if version else "1",
            created_on=(
                proposal.created_at.strftime("%d %b %Y")
                if proposal.created_at
                else ""
            ),
            customer_signature_image_png=_decode_signature_png(
                proposal.customer_signature_image
            ),
            prepared_by_signature_image_png=_decode_signature_png(
                proposal.prepared_by_signature_image
            ),
            prepared_by_name=(
                (
                    prepared_by.get_full_name() or prepared_by.email
                    if prepared_by
                    else ""
                )
                or ""
            ).strip(),
            prepared_by_date=(
                proposal.prepared_by_signed_at.strftime("%d %b %Y")
                if proposal.prepared_by_signed_at
                else ""
            ),
            director_signature_image_png=_decode_signature_png(
                proposal.director_signature_image
            ),
            director_name=(
                (
                    director.get_full_name() or director.email
                    if director
                    else ""
                )
                or ""
            ).strip(),
            director_date=(
                proposal.director_signed_at.strftime("%d %b %Y")
                if proposal.director_signed_at
                else ""
            ),
        )


#: ISO 4217 → currency glyph for the handful of codes we actually
#: quote in. Keeps the rendered PDF readable without pulling a full
#: locale-aware money formatter; unknown codes fall through as a
#: plain ``0.00`` so the number is never lost when someone types an
#: unusual currency.
_CURRENCY_GLYPHS: dict[str, str] = {
    "GBP": "£",
    "EUR": "€",
    "USD": "$",
}


def _format_money(value, currency: str = "GBP") -> str:
    """Render a decimal as a pretty currency string.

    Prefixes the configured glyph (``£`` for GBP) when the currency
    is one we recognise; otherwise leaves the number bare so the
    column header's ``(GBP)`` annotation is still the source of
    truth. ``None`` renders as ``TBC`` — sales prefers that over a
    zero-pound quote showing up on a draft.
    """

    if value is None:
        return "TBC"
    glyph = _CURRENCY_GLYPHS.get(str(currency).upper(), "")
    try:
        number = (
            value
            if isinstance(value, Decimal)
            else Decimal(str(value))
        )
    except Exception:
        return str(value)
    return f"{glyph}{number:,.2f}"


def _decode_signature_png(data_url: str) -> bytes | None:
    """Strip the ``data:image/png;base64,`` prefix and return raw bytes.

    Customers sign on a canvas that emits a PNG data URL — we persist
    it on the proposal. When we stamp the signature into the DOCX we
    need the raw bytes so :mod:`python-docx`'s ``add_picture`` can
    embed them. Any non-PNG prefix or decode error returns ``None``
    and the template keeps its blank ``Signature: ……`` line.
    """

    if not data_url:
        return None
    marker = "base64,"
    idx = data_url.find(marker)
    if idx < 0:
        return None
    try:
        return base64.b64decode(data_url[idx + len(marker) :], validate=True)
    except Exception:
        return None


# ---------------------------------------------------------------------------
# Cell + paragraph helpers — python-docx's ``.text = `` destroys every
# run's formatting, so we write into the first run when possible and
# zero the rest instead of mass-replacing.
# ---------------------------------------------------------------------------


def _set_cell_text(cell: _Cell, value: str) -> None:
    """Replace ``cell``'s text while preserving the first paragraph's
    font / alignment. Extra paragraphs (rarely seen in the templates)
    are flattened into a single line break joined list so we never
    orphan formatting runs that outlive the write."""

    text = value or ""
    paragraphs = cell.paragraphs or []
    if not paragraphs:
        cell.add_paragraph(text)
        return
    first = paragraphs[0]
    if first.runs:
        first.runs[0].text = text
        for extra in first.runs[1:]:
            extra.text = ""
    else:
        first.add_run(text)
    for extra in paragraphs[1:]:
        extra._element.getparent().remove(extra._element)


def _replace_in_paragraph(paragraph, needle: str, replacement: str) -> bool:
    """Swap the first occurrence of ``needle`` in a paragraph's full
    text with ``replacement`` without losing the surrounding runs.

    Implementation: concatenate every run's text, find the match,
    zero every run and write the new string into the first. Loses
    inline mixed formatting inside the needle, but the replaced
    strings are never mixed-formatted in the source templates.
    Returns ``True`` when a replacement happened so the caller can
    stop after the first hit (``Ref:`` and ``Dear ,`` each appear
    once per template)."""

    runs = paragraph.runs
    if not runs:
        return False
    joined = "".join(r.text or "" for r in runs)
    if needle not in joined:
        return False
    new_text = joined.replace(needle, replacement, 1)
    runs[0].text = new_text
    for extra in runs[1:]:
        extra.text = ""
    return True


def _replace_any(paragraph, needles: Iterable[str], replacement: str) -> bool:
    for needle in needles:
        if _replace_in_paragraph(paragraph, needle, replacement):
            return True
    return False


def _paint_customer_table(table: Table, ctx: ProposalRenderContext) -> None:
    """Fill the six-row customer-details table.

    Matches rows by the label string in column A rather than the row
    index so a template tweak (adding / re-ordering rows) doesn't
    silently write into the wrong field.
    """

    values: dict[str, str] = {
        "Name": ctx.customer_name,
        "Email": ctx.customer_email,
        "Phone": ctx.customer_phone,
        "Company Name": ctx.customer_company,
        "Invoice Address": ctx.invoice_address,
        "Delivery Address": ctx.delivery_address,
    }
    for row in table.rows:
        if len(row.cells) < 2:
            continue
        label_raw = row.cells[0].text.strip().rstrip(":").strip()
        value = values.get(label_raw)
        if value is None:
            continue
        _set_cell_text(row.cells[1], value)


def _paint_product_table(
    table: Table, ctx: ProposalRenderContext
) -> None:
    """Fill the pricing table with one row per :class:`ProposalLine`.

    Template layout: the original .docx has a header row plus one
    empty data row. We expand it to ``len(ctx.line_rows)`` rows so a
    three-product proposal prints all three in the same table. Any
    pre-existing empty rows after the header are reused first; new
    rows are appended as needed.
    """

    header_idx = None
    for idx, row in enumerate(table.rows):
        cells = [c.text.strip() for c in row.cells]
        if cells and cells[0].lower().startswith("product code"):
            header_idx = idx
            break
    if header_idx is None:
        return

    row_count_needed = max(len(ctx.line_rows), 1)
    existing_data_rows = len(table.rows) - header_idx - 1
    for _ in range(max(0, row_count_needed - existing_data_rows)):
        table.add_row()

    for line_idx, row_values in enumerate(ctx.line_rows):
        target_row = table.rows[header_idx + 1 + line_idx]
        if len(target_row.cells) < 5:
            continue
        columns = [
            row_values.product_code or "—",
            row_values.description or "—",
            row_values.quantity or "—",
            row_values.unit_price or "TBC",
            row_values.amount or "TBC",
        ]
        for cell, value in zip(target_row.cells[:5], columns):
            _set_cell_text(cell, value)


def _paint_signature_table(
    table: Table, ctx: ProposalRenderContext
) -> None:
    """Replace the dotted ``Date:…`` / ``Signature:…`` placeholders with
    the captured date + signature image (when available)."""

    for row in table.rows:
        if not row.cells:
            continue
        cell = row.cells[0]
        text = cell.text.strip()
        if text.startswith("Date:"):
            date_value = ctx.date_line or "________________________"
            _set_cell_text(cell, f"Date: {date_value}")
        elif text.startswith("Signature:"):
            if ctx.customer_signature_image_png:
                # Clear the cell, drop in the image on a fresh run.
                paragraphs = list(cell.paragraphs)
                for extra in paragraphs[1:]:
                    extra._element.getparent().remove(extra._element)
                if not paragraphs:
                    cell.add_paragraph()
                target = cell.paragraphs[0]
                target.clear()
                target.add_run("Signature: ")
                target.add_run().add_picture(
                    io.BytesIO(ctx.customer_signature_image_png),
                    width=Inches(2.5),
                )
            else:
                _set_cell_text(
                    cell,
                    "Signature: ________________________________",
                )


def _append_internal_signatures(
    doc: _Document, ctx: ProposalRenderContext
) -> None:
    """Stamp the Prepared-by / Director signatures at the end of the
    document so the customer receives a PDF that already carries the
    internal approvals.

    The original Word templates only have a customer signature slot
    near the top. The user explicitly asked for internal sign-off to
    show up on what the client sees, so we append a small two-column
    table after the last paragraph. Signatures render as inline
    images when captured; otherwise the cell shows the role + name +
    date so the customer can see who approved the offer.
    """

    # Skip entirely when neither internal signatory is set — keeps
    # a draft proposal's preview from growing a blank signature
    # footer that would just confuse the scientist.
    has_prepared = bool(ctx.prepared_by_signature_image_png or ctx.prepared_by_name)
    has_director = bool(ctx.director_signature_image_png or ctx.director_name)
    if not (has_prepared or has_director):
        return

    doc.add_paragraph()  # spacing
    heading = doc.add_paragraph()
    run = heading.add_run("Internal approvals")
    run.bold = True

    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    slots = (
        (
            "Prepared by",
            ctx.prepared_by_name,
            ctx.prepared_by_date,
            ctx.prepared_by_signature_image_png,
        ),
        (
            "Approved by (Director)",
            ctx.director_name,
            ctx.director_date,
            ctx.director_signature_image_png,
        ),
    )
    # Start with one row of headers, then add one row per slot so the
    # table fills evenly even when only one slot has a signature.
    header_cells = table.rows[0].cells
    for cell, title in zip(header_cells, ("Prepared by", "Approved by (Director)")):
        _set_cell_text(cell, title)

    body_row = table.add_row()
    for cell, (_, name, date, image_png) in zip(body_row.cells, slots):
        paragraphs = list(cell.paragraphs)
        for extra in paragraphs[1:]:
            extra._element.getparent().remove(extra._element)
        target = cell.paragraphs[0] if paragraphs else cell.add_paragraph()
        target.clear()
        if image_png:
            target.add_run().add_picture(io.BytesIO(image_png), width=Inches(2.0))
        # Name + date under the image (or as the primary content when
        # no signature is captured). Mirrors the DOCX template's
        # ``Name\nDate`` convention below the customer signature.
        cell.add_paragraph(name or "—")
        if date:
            cell.add_paragraph(date)


def _paint_freight_total(
    paragraph, ctx: ProposalRenderContext, needle: str, value: str
) -> bool:
    """Attach ``value`` to a Ready-to-Go-only footer line like
    ``Freight (UK only):`` or ``Total (VAT excluded):``.

    ``value`` already carries the currency glyph (``£1,234.56``) when
    we know the currency, so the ISO code is only appended as a
    disambiguator for currencies we don't have a glyph for — keeps
    the footer tidy instead of reading "£7.14 GBP".
    """

    if needle not in paragraph.text:
        return False
    runs = paragraph.runs
    if not runs:
        return False
    joined = "".join(r.text or "" for r in runs)
    glyph = _CURRENCY_GLYPHS.get((ctx.currency or "").upper(), "")
    if glyph and value.startswith(glyph):
        suffix = f" {value}"
    else:
        suffix = f" {value} {ctx.currency}".rstrip()
    if joined.rstrip().endswith(":"):
        joined = joined.rstrip() + suffix
    else:
        joined = joined.rstrip().rstrip(":") + ":" + suffix
    runs[0].text = joined
    for extra in runs[1:]:
        extra.text = ""
    return True


# ---------------------------------------------------------------------------
# Top-level render functions
# ---------------------------------------------------------------------------


#: XML namespaces the Dynamics quote templates speak. The ``w`` + ``w15``
#: pair is Word 2013+; everything else is unused by us but kept here
#: so the lookup dict is self-documenting.
_DOCX_NS = {
    "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
    "w15": "http://schemas.microsoft.com/office/word/2012/wordml",
}
_W = "{%s}" % _DOCX_NS["w"]
_W15 = "{%s}" % _DOCX_NS["w15"]


def _sdt_binding_xpath(sdt: etree._Element) -> str | None:
    """Return the ``w:xpath`` of a data-bound SDT, or ``None``.

    CRM templates emit ``<w15:dataBinding>`` (Word 2013+). Older
    exports use ``<w:dataBinding>`` — we check both so the same
    filler works against every Dynamics vintage we've seen.
    """

    pr = sdt.find("w:sdtPr", _DOCX_NS)
    if pr is None:
        return None
    for tag in (f"{_W15}dataBinding", f"{_W}dataBinding"):
        binding = pr.find(tag)
        if binding is not None:
            return binding.get(f"{_W}xpath")
    return None


def _classify_xpath(xpath: str) -> tuple[str, str] | None:
    """Reduce a long CRM XPath into a (domain, field) pair.

    The templates use paths like
    ``/ns0:DocumentTemplate[1]/quote[1]/billto_contactname[1]`` or
    ``/ns0:DocumentTemplate[1]/quote[1]/quote_details[1]/productnumber[1]``.
    We only care about the last one or two segments; everything before
    is fixed boilerplate. ``domain`` is either ``"quote"`` (top-level
    proposal field) or ``"line"`` (per-product-row field).
    """

    if not xpath:
        return None
    # Strip ``[1]`` index suffixes and any XML namespace prefix so the
    # match runs against bare field names.
    segments = [
        re.sub(r"\[\d+\]$", "", segment).split(":")[-1]
        for segment in xpath.split("/")
        if segment
    ]
    if not segments:
        return None
    last = segments[-1]
    if "quote_details" in segments:
        return ("line", last)
    if "quote" in segments:
        return ("quote", last)
    return ("quote", last)


def _set_sdt_text(sdt: etree._Element, value: str) -> None:
    """Replace the visible text of a content control.

    Strategy: keep the first ``w:t`` (and its parent run) for its
    formatting, wipe every other ``w:t``, and write the new value
    into the kept one. Handles multi-line values (invoice addresses)
    by inserting ``w:br`` elements between the split lines.
    """

    content = sdt.find("w:sdtContent", _DOCX_NS)
    if content is None:
        return
    text_nodes = content.findall(".//w:t", _DOCX_NS)
    if not text_nodes:
        return

    first = text_nodes[0]
    lines = value.split("\n")
    first.text = lines[0]
    # Clean any leading/trailing whitespace xml:space attribute so
    # Word honours spaces the way we typed them.
    first.set(
        "{http://www.w3.org/XML/1998/namespace}space", "preserve"
    )

    # Zero out remaining text nodes so the combined visible string is
    # just the value we wrote into the first node.
    for extra in text_nodes[1:]:
        extra.text = ""

    # Handle multi-line values (e.g. invoice addresses) by appending
    # <w:br/> line breaks + extra runs after the first text run.
    if len(lines) > 1:
        parent_run = first.getparent()
        if parent_run is None:
            return
        # Siblings following ``parent_run`` in the same paragraph get
        # moved to a later paragraph; simplest is to append our new
        # runs directly after the first text run.
        insert_after = parent_run
        for line in lines[1:]:
            br_run = etree.SubElement(parent_run.getparent(), f"{_W}r")
            etree.SubElement(br_run, f"{_W}br")
            t_run = etree.SubElement(parent_run.getparent(), f"{_W}r")
            t_el = etree.SubElement(t_run, f"{_W}t")
            t_el.text = line
            t_el.set(
                "{http://www.w3.org/XML/1998/namespace}space", "preserve"
            )
            # Move the new runs immediately after insert_after so
            # line 2 renders below line 1 (etree.SubElement appends at
            # the end of parent; we reparent manually).
            parent = parent_run.getparent()
            parent.remove(br_run)
            parent.remove(t_run)
            insert_after.addnext(t_run)
            insert_after.addnext(br_run)
            insert_after = t_run


def _fill_crm_data_bindings(
    docx_bytes: bytes, ctx: ProposalRenderContext
) -> bytes:
    """Second-pass filler for Dynamics-CRM data-bound content controls.

    The original .docx is a zip; ``word/document.xml`` is the layout.
    We parse it, iterate every ``<w:sdt>`` that carries a
    ``<w15:dataBinding>``, and replace the inner text with the real
    value. A ``<w15:repeatingSectionItem>`` wraps the per-product
    row — we clone that item once per :class:`ProposalLine` so a
    multi-product proposal prints every row.
    """

    quote_values = {
        "quotenumber": ctx.reference or "",
        "customeridname": ctx.customer_company or ctx.customer_name or "",
        "billto_contactname": ctx.customer_name or "",
        "billto_telephone": ctx.customer_phone or "",
        # Composite slots carry the raw address only. CRM's default
        # behaviour concatenates contact + company + address but the
        # sales team calls that "name repeated three times" — the
        # top-level customeridname + billto_contactname SDTs already
        # cover the name + company lines above this block.
        "billto_composite": (ctx.invoice_address or "").strip() or "—",
        "shipto_composite": (
            ctx.delivery_address or ctx.invoice_address or ""
        ).strip() or "—",
        "emailaddress": ctx.customer_email or "",
        "description": ctx.product_description or "",
        # Header-band fields. Live in ``word/header2.xml`` on both
        # templates rather than the body, so we have to walk every
        # XML part below — not just document.xml.
        "revisionnumber": ctx.revision_number or "",
        "createdon": ctx.created_on or "",
    }

    # Open the docx, work on a copy of every XML part in memory.
    source = io.BytesIO(docx_bytes)
    with zipfile.ZipFile(source, "r") as zin:
        member_names = zin.namelist()
        payload: dict[str, bytes] = {
            name: zin.read(name) for name in member_names
        }

    # SDT bindings live in the body and also in headers / footers —
    # ``revisionnumber`` / ``createdon`` sit on the page header of
    # both templates (``word/header2.xml``). Process every XML part
    # with bindings so no placeholder slips through.
    bindable_parts = [
        name
        for name in member_names
        if name.startswith(("word/document", "word/header", "word/footer"))
        and name.endswith(".xml")
    ]

    for part_name in bindable_parts:
        raw = payload.get(part_name)
        if raw is None:
            continue
        try:
            root = etree.fromstring(raw)
        except etree.XMLSyntaxError:
            continue

        # Only the body has the repeating line section. Running the
        # expansion on headers / footers is a no-op but cheap.
        _expand_repeating_line_section(root, ctx.line_rows)

        for sdt in root.findall(".//w:sdt", _DOCX_NS):
            xpath = _sdt_binding_xpath(sdt)
            if xpath is None:
                continue
            classified = _classify_xpath(xpath)
            if classified is None:
                continue
            domain, field = classified
            if domain == "quote":
                value = quote_values.get(field)
                if value is None:
                    continue
                _set_sdt_text(sdt, value)
            elif domain == "line":
                pricing_row = _nearest_pricing_row(sdt)
                row_idx = (
                    int(pricing_row.get("proposal_line_idx"))
                    if pricing_row is not None
                    and pricing_row.get("proposal_line_idx") is not None
                    else 0
                )
                if row_idx >= len(ctx.line_rows):
                    continue
                row = ctx.line_rows[row_idx]
                mapping = {
                    "productnumber": row.product_code,
                    "productidname": row.description,
                    "quantity": row.quantity,
                    "priceperunit": row.unit_price,
                    "extendedamount": row.amount,
                }
                value = mapping.get(field)
                if value is None:
                    continue
                _set_sdt_text(sdt, value)

        payload[part_name] = etree.tostring(
            root, xml_declaration=True, encoding="UTF-8", standalone=True
        )

    out = io.BytesIO()
    with zipfile.ZipFile(out, "w", compression=zipfile.ZIP_DEFLATED) as zout:
        for name in member_names:
            zout.writestr(name, payload[name])
    return out.getvalue()


#: Line-level SDT field names — used to pick the table row the
#: Dynamics template laid out for ``quote_details`` so we can clone
#: it once per :class:`ProposalLine`.
_LINE_FIELD_NAMES = frozenset(
    {
        "productnumber",
        "productidname",
        "quantity",
        "priceperunit",
        "extendedamount",
    }
)


def _is_line_sdt(sdt: etree._Element) -> bool:
    """Return ``True`` iff the SDT is bound to a line-level field.

    Used to identify the pricing-row template so the row can be
    cloned per :class:`ProposalLine`. Top-level quote fields
    (customeridname, billto_*, etc.) live outside the pricing table
    and must stay put.
    """

    xpath = _sdt_binding_xpath(sdt)
    if xpath is None:
        return False
    classified = _classify_xpath(xpath)
    return classified is not None and classified[0] == "line"


def _expand_repeating_line_section(
    root: etree._Element, rows: list[ProposalLineRow]
) -> None:
    """Clone the pricing table row once per proposal line.

    The Dynamics templates ship a single ``<w:tr>`` that holds the
    line-level SDTs (``productnumber``, ``productidname``,
    ``quantity``, ``priceperunit``, ``extendedamount``). For a
    three-product proposal we duplicate that row three times so
    LibreOffice renders three priced lines. Every clone (plus the
    original) is tagged with ``proposal_line_idx`` so
    :func:`_fill_crm_data_bindings` knows which line's values to
    write into which row.

    The templates use ``<w:tbl>`` directly rather than CRM's newer
    ``<w15:repeatingSectionItem>``, so we look for rows that contain
    at least one line-level SDT and clone those.
    """

    # Collect the set of rows we need to duplicate. Use a list so
    # duplicate rows (two identical pricing tables in the template)
    # each expand separately.
    pricing_rows: list[etree._Element] = []
    seen_ids: set[int] = set()
    for sdt in root.findall(".//w:sdt", _DOCX_NS):
        if not _is_line_sdt(sdt):
            continue
        row = _ancestor(sdt, f"{_W}tr")
        if row is None or id(row) in seen_ids:
            continue
        seen_ids.add(id(row))
        pricing_rows.append(row)

    for template_row in pricing_rows:
        parent = template_row.getparent()
        if parent is None:
            continue
        template_row.set("proposal_line_idx", "0")
        # Build the clones up-front, then insert them right after the
        # template row so they render in order. ``addnext`` places the
        # new node immediately after the current one, so we iterate
        # reversed so later clones don't leapfrog earlier ones.
        for i in reversed(range(1, max(len(rows), 1))):
            clone = copy.deepcopy(template_row)
            clone.set("proposal_line_idx", str(i))
            template_row.addnext(clone)
        if not rows:
            # No lines at all — strip the placeholder row so the
            # customer doesn't see literal "productnumber" text.
            parent.remove(template_row)


def _ancestor(
    node: etree._Element, tag: str
) -> etree._Element | None:
    """Walk up the tree until an ancestor's tag matches ``tag``."""

    parent = node.getparent()
    while parent is not None:
        if parent.tag == tag:
            return parent
        parent = parent.getparent()
    return None


def _nearest_pricing_row(
    node: etree._Element,
) -> etree._Element | None:
    """Walk up until we hit the cloned ``<w:tr>`` that carries a
    ``proposal_line_idx`` attribute. Used to tell which cloned line
    row an SDT belongs to so we fill the right line's values."""

    parent = node.getparent()
    while parent is not None:
        if (
            parent.tag == f"{_W}tr"
            and parent.get("proposal_line_idx") is not None
        ):
            return parent
        parent = parent.getparent()
    return None


def _pick_template(proposal: Proposal) -> Path:
    template_type = (
        proposal.template_type or ProposalTemplateType.CUSTOM.value
    )
    path = TEMPLATE_PATHS.get(template_type)
    if path is None or not path.exists():
        raise FileNotFoundError(
            f"Proposal template not found for {template_type!r}"
        )
    return path


def render_docx_bytes(proposal: Proposal) -> bytes:
    """Fill the selected .docx template with proposal data and
    return the rendered bytes."""

    template_path = _pick_template(proposal)
    ctx = ProposalRenderContext.from_proposal(proposal)
    doc: _Document = Document(str(template_path))

    freight_text = ctx.freight_amount or "TBC"
    total_text = ctx.total_excl_vat or "TBC"

    # Closing "Yours sincerely, …" block. The templates ship with a
    # hard-coded individual (``Matthew Bowden`` / ``Mathew Bowden``)
    # plus a company line ("Vita Manufacture Sales Team"). We rewrite
    # the individual to the project's assigned sales person, and
    # strip the company line entirely — sales decided the signature
    # should name a real person and nothing else. Transitions to
    # ``in_review`` / ``approved`` / ``sent`` are gated on having a
    # sales person server-side, so in a customer-facing render this
    # is always populated.
    signatory_candidates = ("Matthew Bowden", "Mathew Bowden")
    signatory_value = ctx.sales_person_name or ""
    team_candidates = (
        "Vita Manufacture Sales Team",
        "Vita Manufacture Ltd Sales Team",
        "Vita NPD Sales Team",
    )

    # NB: we no longer touch "Ref:" and "Dear ," here — those lines
    # have CRM SDT data-binding placeholders for ``quotenumber`` /
    # ``customeridname`` / ``billto_contactname`` living inside the
    # same paragraph. Writing into the plain text runs here and then
    # letting the SDT filler fill the placeholders below used to
    # produce double-printed strings like "Ref: PROP-0006 – PRJ-0004
    # - CompanyPROP-0006Company". The SDT filler now owns those
    # placeholders entirely.
    #
    # Collect the team-line paragraphs first so we can detach them
    # from the XML tree after the pass — mutating ``doc.paragraphs``
    # mid-iteration skips elements and risks leaving a stale cursor
    # behind.
    paragraphs_to_remove: list = []
    for p in doc.paragraphs:
        _paint_freight_total(p, ctx, "Freight (UK only):", freight_text)
        _paint_freight_total(
            p, ctx, "Total (VAT excluded):", total_text
        )
        if signatory_value:
            _replace_any(p, signatory_candidates, signatory_value)
        if any(team in p.text for team in team_candidates):
            paragraphs_to_remove.append(p)

    for p in paragraphs_to_remove:
        element = p._element
        parent = element.getparent()
        if parent is not None:
            parent.remove(element)

    # Tables — the originals use the same layout in both templates:
    #   [0] product pricing (re-filled at end of body)
    #   [1] customer details
    #   [2] confirmation check-boxes (left intact)
    #   [3] Date + Signature
    #   [5] phases (static copy, leave alone)
    #   [6] product pricing duplicate
    # The Dynamics templates include the Product Code pricing table
    # twice — once up top under the "commercial price" intro and
    # once at the very bottom. Both bind to the same
    # ``quote_details``, so a one-line proposal ends up with two
    # identical pricing rows in the PDF. R&D / sales consistently
    # read that as a bug; we keep only the first occurrence.
    pricing_table_seen = False
    for idx, table in enumerate(doc.tables):
        header_cells = [c.text.strip() for c in table.rows[0].cells] if table.rows else []
        if header_cells and header_cells[0].lower().startswith("product code"):
            # Skip painting here — the pricing table content is
            # driven by the CRM SDT data-binding pass below. But
            # *do* strip any duplicate occurrence so the customer
            # only sees one pricing block.
            if pricing_table_seen:
                table_element = table._element
                parent = table_element.getparent()
                if parent is not None:
                    parent.remove(table_element)
            pricing_table_seen = True
            continue
        elif any("Name:" in c.text for c in (table.rows[0].cells if table.rows else [])):
            _paint_customer_table(table, ctx)
        elif any(
            "Date:" in c.text or "Signature:" in c.text
            for row in table.rows
            for c in row.cells
        ):
            _paint_signature_table(table, ctx)

    # Append the Prepared-by / Director block so the customer PDF
    # shows who inside Vita NPD already approved the offer.
    _append_internal_signatures(doc, ctx)

    buffer = io.BytesIO()
    doc.save(buffer)

    # Second pass — the templates were exported from Dynamics 365 CRM
    # and every dynamic field (customer name, product number, prices)
    # lives inside ``<w:sdt>`` content controls bound via
    # ``<w15:dataBinding>`` to an internal XML data store. python-docx
    # treats those as opaque blobs, so python-docx's Find & Replace
    # never touched them. We post-process the saved .docx through
    # lxml to fill each bound SDT with the real value, preserving the
    # run properties (font, colour, bold) the template was styled in.
    return _fill_crm_data_bindings(buffer.getvalue(), ctx)


def _render_pdf_bytes_uncached(proposal: Proposal) -> bytes | None:
    """Actual DOCX → PDF pipeline. Slow (5-15s per run) because it
    drives LibreOffice. Call through :func:`render_pdf_bytes` so
    repeated requests hit the cache.

    Word is intentionally *not* a fallback here. On a macOS dev
    machine ``docx2pdf`` opens Microsoft Word visibly and blocks
    for minutes when it hits a license prompt or a disabled
    AppleScript permission — the user can see the Word app pop
    up while they're trying to view a proposal. LibreOffice
    ``soffice --headless`` is silent + deterministic, so we rely
    on it exclusively. If soffice isn't installed we return
    ``None`` and the HTTP handler falls back to the HTML render.
    """

    docx_bytes = render_docx_bytes(proposal)

    # Serialise across concurrent renders — see note on
    # ``_RENDER_LOCK``. Requests still arrive in parallel from the
    # client but the CPU-heavy conversion only runs one at a time,
    # which cuts tail latency on a shared dev machine.
    with _RENDER_LOCK:
        with tempfile.TemporaryDirectory() as tmp_dir:
            tmp_path = Path(tmp_dir)
            docx_path = tmp_path / "proposal.docx"
            docx_path.write_bytes(docx_bytes)
            pdf_path = tmp_path / "proposal.pdf"

            if _try_libreoffice(docx_path, tmp_path) and pdf_path.exists():
                return pdf_path.read_bytes()
    return None


def _proposal_digest(proposal: Proposal) -> str:
    """Compact fingerprint of the proposal's render-relevant state.

    Used as the cache key for :attr:`Proposal.rendered_pdf`. When the
    digest matches the stored value the cached PDF is still valid;
    otherwise we regenerate. Includes every line's ``updated_at``
    and the proposal's own ``updated_at`` so edits at any level
    invalidate the cache.

    The sales person's name is ``get_full_name()`` dependent so it
    appears on the rendered PDF; changes to the user's profile
    aren't reflected until the proposal itself is touched. Acceptable
    trade-off — name edits on sales accounts are rare and the
    alternative (hash the user row too) adds an FK round-trip on
    every cache-check.
    """

    parts = [
        f"v{_RENDER_VERSION}",
        str(proposal.updated_at.timestamp()),
        str(proposal.status),
        # Sales person override is rendered into the signature line.
        # ``updated_at`` already moves when the override is set via
        # ``update_proposal``, so including the id is belt-and-braces
        # — covers the case where the underlying user's name changes
        # (``get_full_name``) without the proposal itself being
        # touched.
        str(proposal.sales_person_id or ""),
    ]
    for line in proposal.lines.order_by("id").values_list(
        "id", "updated_at"
    ):
        parts.append(f"{line[0]}:{line[1].timestamp()}")
    return "|".join(parts)


def render_pdf_bytes(proposal: Proposal) -> bytes | None:
    """Return the PDF bytes for this proposal, caching between calls.

    Rationale — each conversion takes several seconds because it
    drives Microsoft Word via AppleScript (or LibreOffice headless).
    Without a cache, every page load + every Download click queues a
    fresh conversion and the Word process can pile up. The cache
    key is a cheap digest of the proposal's ``updated_at`` plus
    every line's ``updated_at``; any edit invalidates the cached
    bytes so the next render regenerates.

    Returns ``None`` when no converter is available — the HTTP layer
    falls back to the HTML template for that call.
    """

    digest = _proposal_digest(proposal)
    if (
        proposal.rendered_pdf_digest == digest
        and proposal.rendered_pdf
    ):
        return bytes(proposal.rendered_pdf)

    pdf_bytes = _render_pdf_bytes_uncached(proposal)
    if pdf_bytes is None:
        return None

    # Persist the new cache atomically. ``update_fields`` avoids
    # touching ``updated_at`` — which would instantly invalidate the
    # digest we just wrote.
    proposal.rendered_pdf = pdf_bytes
    proposal.rendered_pdf_digest = digest
    type(proposal).objects.filter(pk=proposal.pk).update(
        rendered_pdf=pdf_bytes,
        rendered_pdf_digest=digest,
    )
    return pdf_bytes


def _try_libreoffice(docx_path: Path, out_dir: Path) -> bool:
    """Best-effort LibreOffice headless conversion.

    Returns ``True`` iff soffice successfully wrote a PDF next to
    ``docx_path`` in ``out_dir``. Silent failure surfaces as ``False``
    so the caller can fall back.

    LibreOffice locks a single user profile directory by default, so
    two parallel calls (e.g. two browser tabs loading the same
    proposal) crash the second with ``UserInstallation already in
    use``. We give every call its own profile under ``out_dir`` via
    ``-env:UserInstallation=file://…`` — cheap to create, thrown
    away with the temp dir.
    """

    candidates = ["soffice", "libreoffice"]
    profile_dir = out_dir / "lo-profile"
    profile_uri = f"file://{profile_dir}"
    for name in candidates:
        exe = shutil.which(name)
        if exe is None:
            continue
        cmd = [
            exe,
            f"-env:UserInstallation={profile_uri}",
            "--headless",
            "--nologo",
            "--nofirststartwizard",
            "--norestore",
            "--nolockcheck",
            "--convert-to",
            "pdf",
            "--outdir",
            str(out_dir),
            str(docx_path),
        ]
        try:
            result = subprocess.run(
                cmd,
                capture_output=True,
                timeout=120,
            )
        except Exception:
            continue
        # ``soffice`` returns 0 even when conversion silently fails
        # (e.g. missing template). Verify the PDF landed where we
        # expect before declaring success.
        expected_pdf = out_dir / f"{docx_path.stem}.pdf"
        if result.returncode == 0 and expected_pdf.exists():
            return True
    return False


def _try_word(docx_path: Path, pdf_path: Path) -> bool:
    """Convert via Microsoft Word on macOS. Word's AppleScript
    interface is the only way to get a pixel-faithful PDF out of
    .docx without LibreOffice. Silently returns ``False`` when Word
    isn't available or the conversion throws."""

    try:
        # Lazy import so servers without ``docx2pdf`` (Linux /
        # containers) don't crash at import time.
        from docx2pdf import convert  # type: ignore[import-not-found]
    except Exception:
        return False
    try:
        # ``docx2pdf.convert(src, dst)`` accepts explicit output paths
        # when both are files — skips the Word "batch folder" codepath
        # that sometimes hangs on single-document input.
        convert(str(docx_path), str(pdf_path))
    except Exception:
        return False
    return pdf_path.exists()
